from docx import Document

from .models import CommissionMember


def generate_act(apartment, act):
    doc = Document()

    doc.add_heading('АКТ ИНВЕНТАРИЗАЦИИ ЖИЛОГО ПОМЕЩЕНИЯ', level=1)

    doc.add_paragraph(f'Адрес: {apartment.house.address}')
    doc.add_paragraph(f'Квартира № {apartment.number}')
    doc.add_paragraph(f'Площадь: {apartment.total_area} м²')

    doc.add_heading('Техническое состояние')
    doc.add_paragraph(act.technical_condition)

    doc.add_heading('Комиссия')
    commission = CommissionMember.objects.filter(year=act.year)

    for member in commission:
        doc.add_paragraph(f'{member.full_name} — {member.position}')

    file_path = f'media/acts/act_{act.id}.docx'
    doc.save(file_path)

    return file_path
